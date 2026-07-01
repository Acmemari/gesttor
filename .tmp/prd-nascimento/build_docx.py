# -*- coding: utf-8 -*-
"""
Gera o documento Word (.docx) com o PRD da tela de Nascimento.
Inclui as 8 imagens (mockups) geradas por build.py.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
IMG = ROOT / "img"
OUT = Path("C:/gesttor/.maps") if Path("C:/gesttor/.maps").exists() else ROOT
OUT_FILE = ROOT / "PRD-Tela-Nascimento-Inttegra.docx"

doc = Document()

# ── Página em A4 paisagem para acomodar os mockups largos ────────────────────
for section in doc.sections:
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)

# ── Estilos básicos ──────────────────────────────────────────────────────────
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p.paragraph_format.space_after = Pt(4)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def h4(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    return p

def para(text, *, bold=False, gray=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if gray:
        r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    p.paragraph_format.space_after = Pt(2)
    return p

def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(11)

def num_list(items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it)
        r.font.size = Pt(11)

def hr_para():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "16A34A")
    pBdr.append(bot)
    pPr.append(pBdr)

def add_image(name, caption=None, width_cm=24.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(IMG / name), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = cap.add_run(caption)
        c.italic = True
        c.font.size = Pt(9.5)
        c.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        cap.paragraph_format.space_after = Pt(6)

def add_table(rows, *, header=True, col_widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    if col_widths:
        for c, ww in enumerate(col_widths):
            for r in range(len(rows)):
                table.cell(r, c).width = Cm(ww)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = ""
            par = cell.paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(10)
            if header and r == 0:
                run.bold = True
    return table

# ───────────────────────── CAPA ─────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("INTTEGRA · Pecuário")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x16,0xA3,0x4A); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = p.add_run("PRD — Tela de Nascimentos")
rr.bold = True; rr.font.size = Pt(32); rr.font.color.rgb = RGBColor(0x0F,0x17,0x2A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Movimentação › Nascimento — Especificação completa para o time de Engenharia")
rs.font.size = Pt(13); rs.font.color.rgb = RGBColor(0x6B,0x72,0x80); rs.italic = True

doc.add_paragraph()
meta_tbl = doc.add_table(rows=4, cols=2)
meta_tbl.style = "Light List Accent 1"
meta_pairs = [
    ("Produto",            "INTTEGRA — Sistema de Gestão Pecuária"),
    ("Módulo",             "Pecuário › Movimentação › Nascimento"),
    ("Versão do PRD",      "1.0  ·  22/06/2026"),
    ("Status",             "Em produção (referência para refatoração/reescrita)"),
]
for i, (k, v) in enumerate(meta_pairs):
    c0 = meta_tbl.cell(i, 0); c1 = meta_tbl.cell(i, 1)
    c0.text = ""; c1.text = ""
    r0 = c0.paragraphs[0].add_run(k); r0.bold = True; r0.font.size = Pt(10.5)
    r1 = c1.paragraphs[0].add_run(v); r1.font.size = Pt(10.5)

doc.add_paragraph()
hr_para()
para("Este documento descreve, em detalhe operacional e técnico, a tela de Nascimentos "
     "do Inttegra Pecuário. O objetivo é permitir que o time de desenvolvimento "
     "reconstrua a funcionalidade — em qualquer stack — preservando comportamento, "
     "regras de negócio, modelo de dados e UX. Cada seção combina descrição funcional, "
     "regras de negócio, contratos de API e referência visual (mockup).",
     gray=True, italic=True)

doc.add_page_break()

# ───────────────────────── 1. VISÃO ─────────────────────────
h1("1. Visão e contexto")
para("A tela de Nascimento é o ponto de entrada do principal evento positivo do rebanho "
     "de cria: o nascimento de bezerros. Ela é o lançamento que abastece a Ficha "
     "Animal (categoria atual) e mantém o saldo correto do estoque por categoria. "
     "Foi construída sob o conceito de DUPLA CAMADA ADITIVA (coletivo + individual), "
     "que reflete a realidade do campo: o capataz raramente identifica todos os "
     "bezerros no momento do parto.")

h3("1.1 Objetivos do produto")
bullets([
    "Reduzir atrito no lançamento de nascimentos em massa — fazendas com 200+ partos/mês.",
    "Permitir registrar primeiro a quantidade total (coletivo) e detalhar os animais depois (atribuição de ID), sem perder a base de conciliação.",
    "Unificar a experiência com as demais movimentações (Compra, Venda, Morte) via o kit compartilhado 'Defina seus campos'.",
    "Manter a Ficha Animal sempre coerente: cada nascimento detalhado gera/atualiza a categoria viva do animal e alimenta o cadastro pelo apelido (ID Manejo).",
])

h3("1.2 Público-alvo")
bullets([
    "Capataz/gerente de retiro — lança nascimentos diariamente no celular ou tablet.",
    "Técnico/zootecnista — atribui IDs e completa fichas no escritório.",
    "Diretor/proprietário — usa a aba Registros para auditoria e fechamento de safra.",
])

h3("1.3 Princípios de design")
bullets([
    "Quantidade é a âncora. Tudo é reconciliado contra a quantidade total — categoria é distribuição, não chave.",
    "Coletivo e individual coexistem. Eles SOMAM; nunca substituem um ao outro.",
    "Configurável por organização. Cada fazenda escolhe onde cada campo aparece (kit 'Defina seus campos').",
    "Mesma experiência em massa. Planilha modelo exportável + importação validada para lançamentos em lote.",
    "Mobile-friendly. Painéis fluidos com container queries; tela cheia para foco no campo.",
])

doc.add_page_break()

# ───────────────────────── 2. CONCEITUAÇÃO ─────────────────────────
h1("2. Conceito-chave: Camada Dupla Aditiva")
para("Em uma única tela, há DOIS caminhos para registrar nascimentos. Os dois caminhos "
     "coexistem dentro do MESMO lançamento e a soma deles forma o total. Esse conceito "
     "é o coração das regras de negócio do módulo.")

table = doc.add_table(rows=4, cols=3)
table.style = "Light Grid Accent 1"
header = [("Aspecto", "Modo Coletivo (Lote)", "Modo Individual (Brinco)")]
data = [
    ("Quando usar",    "Não há tempo/condições de individualizar.", "Bezerro já está identificado (brinco/RFID)."),
    ("O que se grava", "Quantidade + categoria (sem ID).",          "Uma FICHA por animal (ID Manejo + atributos)."),
    ("Saída no banco", "campo `naoIdentificados` += qtd.",         "linhas em nascimento_fichas."),
]
for c, val in enumerate(header[0]):
    cell = table.cell(0, c)
    cell.text = ""
    rr = cell.paragraphs[0].add_run(val); rr.bold = True; rr.font.size = Pt(10.5)
for r, row in enumerate(data, start=1):
    for c, val in enumerate(row):
        cell = table.cell(r, c)
        cell.text = ""
        rr = cell.paragraphs[0].add_run(val); rr.font.size = Pt(10.5)

doc.add_paragraph()
para("Estes dois caminhos somam (camada dupla aditiva). Toda categoria mostrada na "
     "tela tem dois números: 'Sem ID' (declarado) e 'Com ID' (detalhado). A soma é o "
     "total da categoria, e o total geral do lançamento é a base de conciliação.",
     gray=True, italic=True)

h3("2.1 Invariantes")
bullets([
    "qtd (total)  =  naoIdentificados + fichas.length",
    "status = 'conciliado'  se  naoIdentificados == 0  ;  senão 'pendente'",
    "Salvar com fichas: subtrai da parte declarada do mesmo catId para não duplicar.",
    "Add ficha (atribuição posterior): naoIdentificados -= 1 e status é recalculado no servidor.",
    "Edição (PUT): substitui as fichas integralmente pelo conjunto reenviado.",
])

doc.add_page_break()

# ───────────────────────── 3. FLUXOS PRINCIPAIS ─────────────────────────
h1("3. Fluxos principais")

# ── 3.1 Modo coletivo ──
h2("3.1 Lançar nascimentos em LOTE (modo coletivo)")
add_image("01_tela_coletivo.png",
          "Figura 1 — Tela inicial, modo coletivo. O toggle 'Lote' (2º ícone) está ativo. "
          "Quantidade é a âncora; +mais permite distribuir em várias categorias.")

h3("Passo a passo")
num_list([
    "Selecionar Data (a Safra é derivada automaticamente: jul→jun; ex.: 22/06/2026 → 2025/2026).",
    "Confirmar Proprietário, Fazenda, Retiro e Local (Local depende do Retiro; quando há um único Retiro, ele é pré-selecionado).",
    "Toggle 'Lote' (visão coletiva — padrão).",
    "Informar Quantidade (cab.) e a Categoria.",
    "(Opcional) Clicar '+ mais' para adicionar outra distribuição de categoria no mesmo lançamento.",
    "(Opcional) Acionar 'Sanitário' para registrar aplicações no nível do lançamento.",
    "Conferir 'Distribuição por categoria' (painel direito) — soma das categorias = Quantidade.",
    "Clicar Salvar. O lançamento entra como 'pendente' enquanto houver animais sem ID.",
])

h4("Regras de negócio")
bullets([
    "A categoria 'sem detalhe' aceita SOMENTE categorias do grupo 'bezerros_mamando'.",
    "Quantidade é livre (inteiro > 0). Não há tetos por dia/lote.",
    "Sem Data ou sem (Quantidade>0 + Categoria) o botão Salvar fica desabilitado.",
    "Cancelar limpa o formulário SEM persistir nada e desativa edição em curso.",
])

doc.add_page_break()

# ── 3.2 Modo individual ──
h2("3.2 Lançar nascimentos COM IDENTIFICAÇÃO (modo individual)")
add_image("02_modo_individual.png",
          "Figura 2 — Modo individual ativo. Os campos coletivos ficam cinza (não-editáveis) "
          "e a tabela 'Defina seus campos' aparece abaixo para incluir bezerros um a um.")

h3("Passo a passo")
num_list([
    "Após preencher o cabeçalho, clicar no toggle 'Brinco' (1º ícone).",
    "O painel 'Defina seus campos' aparece com a linha 'Repete em todos' (data/raça/lote/sanitário) e a linha de inclusão.",
    "Preencher ID Manejo (obrigatório), Categoria (obrigatório), Sexo (deriva da categoria), Porte, Peso etc.",
    "Clicar '+ Adicionar'. O bezerro é incluído na tabela e o formulário sugere o próximo ID Manejo (autoinc opcional — preserva prefixo, sufixo e zeros à esquerda).",
    "Repetir para os demais animais.",
    "Salvar. O lançamento já entra 'conciliado' se todos os bezerros foram identificados.",
])

h4("Importação em massa via planilha")
bullets([
    "Botão 'Planilha ↓' baixa um modelo .xlsx com as colunas EXATAS configuradas para a organização.",
    "Botão 'Planilha ↑' abre um modal que valida o arquivo carregado (linha-a-linha) e mostra OK / Erro / Duplicado.",
    "Linhas válidas viram animais na tabela de detalhe (ainda não persistidas — confirmar Salvar).",
])

doc.add_page_break()

# ── 3.3 Configurar campos ──
h2("3.3 Configurar 'Defina seus campos'")
add_image("03_config_campos.png",
          "Figura 3 — Modal 'Configurar campos do Lançamento Rápido'. Cada campo pode ir para "
          "Linha Superior (repete), Tabela Lançamento (por animal), Dados Adicionais (recolhido) "
          "ou Desativado. Arraste a alça ⠿ para reordenar.")

h3("Regras do modal")
bullets([
    "Campos com 🔒 (ID Manejo) ficam SEMPRE na Tabela — só admite reordenação.",
    "Campos de seção (Sanitário) só admitem Superior ou Desativado.",
    "Toggle 'Numeração automática' (autonum) controla a sugestão do próximo ID Manejo ao adicionar.",
    "A configuração é por ORGANIZAÇÃO, persistida em nascimento_field_configs (1 linha por org).",
    "Botão 'Restaurar padrão' volta todos os places ao default do registry (def: 'top'|'bottom'|'dados').",
])

h4("Catálogo de campos (LR_REGISTRY)")
para("Os campos abaixo formam o registry compartilhado da tela. Cada um declara: id, label, "
     "tipo (text/select/date/cat/sexo/lote/weight/textarea/sanitario), obrigatoriedade (req), "
     "lista de opções (options), placeholder e destino padrão (def).", gray=True, italic=True)

reg_rows = [
    ["id", "label", "tipo", "obrig.", "default", "destino padrão"],
    ["apelido (🔒)", "ID Manejo", "text", "✔", "—", "bottom (tabela)"],
    ["categoria", "Categoria", "cat", "✔", "—", "bottom (tabela)"],
    ["data", "Data", "date", "✔", "hoje", "top (superior)"],
    ["raca", "Raça", "select", "✔", "Nelore", "top"],
    ["lote", "Lote", "lote", "—", "—", "top"],
    ["rfid", "ID Eletrônica", "text", "—", "—", "bottom"],
    ["sisbov", "Nº SISBOV", "text", "—", "—", "bottom"],
    ["sexo", "Sexo", "sexo", "✔", "deriva da categoria", "bottom"],
    ["porte", "Porte", "select", "✔", "M", "bottom"],
    ["colostro", "Colostro?", "select", "—", "Sim", "bottom"],
    ["peso", "Peso nasc.", "weight", "—", "—", "bottom"],
    ["pesagem", "Pesagem", "select", "—", "Manual", "bottom"],
    ["sanitario", "Sanitário", "sanitario", "—", "—", "top (enableOnly)"],
    ["nome", "Nome Completo", "text", "—", "—", "dados (span 2)"],
    ["pesoNascer", "Peso ao Nascer", "weight", "—", "—", "dados"],
    ["grau", "Grau de Sangue", "select", "—", "—", "dados"],
    ["rgn", "RGN/Tatuagem", "text", "—", "—", "dados"],
    ["pelagem", "Pelagem", "select", "—", "—", "dados"],
    ["chifre", "Tipo de Chifre", "select", "—", "—", "dados"],
    ["rgd", "RGD", "text", "—", "—", "dados"],
    ["serie", "Série Alfa", "text", "—", "—", "dados"],
    ["pai", "Pai - ID Usual", "text", "—", "—", "dados"],
    ["mae", "Mãe - ID Usual", "text", "—", "—", "dados"],
    ["obs", "Observação", "textarea", "—", "—", "dados (span 3)"],
]
add_table(reg_rows, col_widths=[3.2, 4.0, 2.6, 1.8, 4.4, 5.0])

para("Além desses, a tela mescla CAMPOS PERSONALIZADOS do cadastro 'Campos Personalizados' "
     "(módulo Pecuário) com ids no formato cp_<uuid>; seus valores ficam em fichas.extras (jsonb).",
     gray=True, italic=True)

doc.add_page_break()

# ── 3.4 Sanitário ──
h2("3.4 Aplicações sanitárias (nível movimento)")
add_image("04_sanitario.png",
          "Figura 4 — Seção 'Sanitário'. Cada aplicação aceita Vacina/Medicamento, "
          "Unidade, Tipo de Dose, Dose e (opcional) Por Cada (X) Kg; o Custo é "
          "calculado automaticamente (custoUnit × dose).")

bullets([
    "A seção é recolhível, acionada por um botão na linha 'Repete em todos'.",
    "Tipo de Aplicação tem 2 modos: Única (livre) ou Protocolo (pré-configurado).",
    "Os medicamentos e protocolos vêm hoje de listas estáticas (MEDICAMENTOS, PROTOCOLOS) — a planejada migração para tabelas vivas está fora do escopo desta tela.",
    "Os itens são salvos como snapshot no jsonb sanitario do movimento (não vinculados por FK).",
])

doc.add_page_break()

# ── 3.5 Atribuir ID ──
h2("3.5 Atribuir ID (individualização posterior)")
add_image("05_atribuir_id.png",
          "Figura 5 — Painel 'Atribuição de ID' acionado a partir do menu ••• do lançamento. "
          "Cada linha adicionada decrementa naoIdentificados e abate a parte declarada da categoria.")

h3("Passo a passo")
num_list([
    "Na aba Registros, clicar ••• do lançamento desejado e escolher 'Atribuir ID'.",
    "O usuário é levado para a aba Lançamento, onde o painel é exibido abaixo do cartão principal.",
    "Preencher os campos da linha de inclusão (ID Manejo, Categoria, RFID, SISBOV, Peso, Porte) — Enter envia.",
    "A cada Adicionar, o painel mostra o resumo atualizado (X de Y detalhados) e a tabela de fichas.",
    "Quando naoIdentificados chegar a 0, o lançamento muda para 'conciliado' automaticamente.",
])

h4("Regras")
bullets([
    "A categoria padrão da linha de inclusão é o 1º catDecl do movimento (ou a 1ª categoria disponível).",
    "ID Manejo é obrigatório e DEVE ser único dentro do escopo do lançamento.",
    "O sexo é derivado da categoria (m/f) — o usuário não escolhe.",
    "Toda inclusão de ficha posterior faz POST com action='add-ficha' (não recria o movimento).",
])

doc.add_page_break()

# ── 3.6 Registros ──
h2("3.6 Aba 'Registros' (histórico master-detail)")
add_image("06_registros.png",
          "Figura 6 — Aba Registros. Master (lista de lançamentos) em cima, Detail (fichas individuais) "
          "embaixo, com régua arrastável. Expanda a linha (▾) para ver e selecionar uma categoria do lançamento.")

h3("Interações")
bullets([
    "Clicar 1× na linha — seleciona e abre o Detalhe no painel inferior.",
    "Clicar 2× (ou na seta ▸) — expande as categorias do lançamento abaixo da linha.",
    "Clicar numa categoria expandida — filtra o Detalhe para mostrar só fichas dessa categoria.",
    "Régua arrastável entre Master e Detail (15%–85%) — útil ao atribuir IDs.",
    "Menu ••• por linha (e por categoria expandida): Ver · Editar · Atribuir ID · Excluir.",
    "Excluir mostra confirmação nativa (DD/MM/YYYY, total). Editar reabre o lançamento no formulário superior.",
])

doc.add_page_break()

# ── 3.7 Planilha ──
h2("3.7 Importação por planilha (lançamento em massa)")
add_image("07_importar_planilha.png",
          "Figura 7 — Modal de importação. Valida linha-a-linha contra o catálogo da organização: "
          "categorias, raças, sexos e duplicidade do ID Manejo (escopo do lançamento aberto).")

bullets([
    "O modelo .xlsx é gerado dinamicamente a partir da CONFIGURAÇÃO atual de campos (ordem/destinos).",
    "Aceita .xlsx, .xls e .csv. Tem que haver linhas além do cabeçalho.",
    "Erros suportados: campo obrigatório vazio, categoria/raça inexistente, ID Manejo duplicado, peso não numérico.",
    "Linhas marcadas como OK podem ser confirmadas em lote — entram como animais detalhados (ainda não persistidas).",
    "O usuário ainda precisa clicar Salvar para persistir o lançamento no banco.",
])

doc.add_page_break()

# ───────────────────────── 4. MODELO DE DADOS ─────────────────────────
h1("4. Modelo de dados")
add_image("08_fluxo_dados.png",
          "Figura 8 — Modelo conceitual. nascimento_movimentos (1) ↔ (N) nascimento_fichas; "
          "ambas referenciam animal_categories. nascimento_field_configs é singleton por organização.")

h3("4.1 Tabelas (Drizzle ORM · PostgreSQL/Neon)")

para("Esquema simplificado, formatado para leitura:", gray=True, italic=True)

h4("nascimento_movimentos (1 linha por LANÇAMENTO)")
schema_mov = [
    ["coluna", "tipo", "regra"],
    ["id",                   "uuid PK",          "defaultRandom"],
    ["organization_id",      "uuid FK",          "→ organizations (cascade)"],
    ["farm_id",              "text FK",          "→ farms (set null)"],
    ["local_id",             "uuid FK",          "→ farm_locais (set null)"],
    ["proprietario_id",      "uuid FK",          "→ people (set null)"],
    ["data",                 "date NOT NULL",    "ISO YYYY-MM-DD"],
    ["safra",                "text",             "derivada por safraDaData()"],
    ["retiro",               "text",             "nome (não-FK)"],
    ["qtd",                  "int NOT NULL",     "default 0"],
    ["nao_identificados",    "int NOT NULL",     "default 0"],
    ["status",               "text NOT NULL",    "'pendente' | 'conciliado'"],
    ["cat_decl",             "jsonb",            "[{catId, qtd}]"],
    ["sanitario",            "jsonb",            "SanItem[] (snapshot)"],
    ["criado_por",           "text FK",          "→ user_profiles (set null)"],
    ["created_at/updated_at","timestamp",        "padrão Drizzle"],
]
add_table(schema_mov, col_widths=[4.5, 4.5, 9.5])

h4("nascimento_fichas (N linhas por movimento, 1 por bezerro identificado)")
schema_fic = [
    ["coluna", "tipo", "regra"],
    ["id",            "uuid PK",          "defaultRandom"],
    ["movimento_id",  "uuid FK NOT NULL", "→ nascimento_movimentos (cascade)"],
    ["categoria_id",  "uuid FK",          "→ animal_categories (set null)"],
    ["apelido",       "text NOT NULL",    "ID Manejo"],
    ["rfid",          "text",             "RFID/Brinco eletrônico"],
    ["sisbov",        "text",             "Nº SISBOV"],
    ["porte",         "text",             "P|M|G"],
    ["raca",          "text",             "—"],
    ["peso",          "numeric(8,2)",     "kg"],
    ["extras",        "jsonb NOT NULL",   "default {} — Campos Personalizados (cp_*)"],
    ["created_at",    "timestamp",        "padrão"],
]
add_table(schema_fic, col_widths=[4.5, 4.5, 9.5])

h4("nascimento_field_configs (singleton por organização)")
schema_cfg = [
    ["coluna",            "tipo",    "regra"],
    ["organization_id",   "uuid",    "unique"],
    ["config",            "jsonb",   "{ places, order, autonum }"],
]
add_table(schema_cfg, col_widths=[5.5, 4.0, 9.0])

doc.add_page_break()

# ───────────────────────── 5. API ─────────────────────────
h1("5. Contratos de API")
para("Endpoint único: /api/nascimentos — autenticado via better-auth (cookie/bearer). "
     "Todos os payloads vão/voltam em JSON com envelope { data, ... } ou { error, code }.",
     gray=True, italic=True)

h2("5.1 GET /api/nascimentos")
bullets([
    "?organizationId=<uuid>  →  lista movimentos da organização (com fichas).",
    "?id=<uuid>             →  busca um movimento (com fichas).",
])

h2("5.2 POST /api/nascimentos — criar movimento")
para("Body:", bold=True)
para("{ organizationId, farmId?, localId?, proprietarioId?, data: 'YYYY-MM-DD', safra?, "
     "retiro?, qtd, naoIdentificados, status, catDecl: [{catId,qtd}], sanitario: [], "
     "fichas: [{ apelido, catId?, rfid?, sisbov?, porte?, raca?, peso?, extras? }] }")
bullets([
    "Validações: organizationId e data obrigatórios; data em formato YYYY-MM-DD; status ∈ {pendente, conciliado}.",
    "Servidor mantém: criado_por = userId; localId = resolveDefaultLocalId(farmId) se omitido.",
])

h2("5.3 POST /api/nascimentos — adicionar ficha (atribuir ID)")
para("Body: { action: 'add-ficha', movimentoId, apelido, categoriaId?, rfid?, sisbov?, porte?, raca?, peso? }")
bullets([
    "Decrementa naoIdentificados (não negativo) e recalcula status.",
    "Retorna o movimento atualizado com a lista completa de fichas.",
])

h2("5.4 PUT /api/nascimentos?id=<uuid> — atualizar")
bullets([
    "Body idêntico ao POST de criação (sem organizationId, sem criado_por).",
    "Substitui as fichas INTEGRALMENTE pelo conjunto reenviado.",
    "Atualiza updated_at; recalcula status pelo naoIdentificados informado.",
])

h2("5.5 DELETE /api/nascimentos?id=<uuid>")
para("Remove o movimento (cascade nas fichas). Retorna { deleted: true }.")

h2("5.6 GET/PUT /api/nascimento-field-config")
bullets([
    "GET ?organizationId=  retorna { places, order, autonum } ou defaults se nada salvo.",
    "PUT body { organizationId, config } — upsert único por organização.",
])

doc.add_page_break()

# ───────────────────────── 6. UI/UX ─────────────────────────
h1("6. UI/UX — Especificação visual")

h3("6.1 Layout do cartão principal (Lançamentos)")
bullets([
    "Cartão único dividido por uma RÉGUA INTERNA: esquerda 65% (formulário), direita 35% (Distribuição).",
    "Container query em 1180px: abaixo disso empilha (régua vira horizontal).",
    "Painéis fluidos preenchem 100% da área (reage ao recolher do sidebar md:ml-64 → md:ml-16).",
])

h3("6.2 Tipografia e cores")
bullets([
    "Fonte: stack do sistema (Inter como base).",
    "Verde primário: #16a34a · Verde escuro (hover): #15803d · Verde claro (chips): #e7f6ec.",
    "Laranja de aviso: #ea580c · Vermelho de erro: #dc2626 · Azul de identificado: #2563eb.",
    "Borda padrão: #e5e7eb · Fundo do app: #f9fafb · Texto principal: #0F172A.",
])

h3("6.3 Microinterações")
bullets([
    "Toggle Brinco/Lote: o ÚLTIMO clicado fica ativo (anel verde, fundo claro).",
    "Quando há 1 único retiro na fazenda, ele já vem selecionado (não pergunta).",
    "Resumo ao vivo: 'Total X cab. · Y identificados · Z a detalhar' (chip cinza quando há pendência, verde quando 100% conciliado).",
    "Tela cheia: ícone de Expandir no painel 'Defina seus campos' — Esc fecha; trava o scroll do fundo.",
    "Próximo ID Manejo: sugerido ao Adicionar, mantendo prefixo/sufixo/zeros (504A → 505A · BZ-09 → BZ-10).",
])

h3("6.4 Acessibilidade")
bullets([
    "Foco visível em todos os controles (ring verde 3px com transparência).",
    "Régua arrastável tem role='separator' e aria-label próprio.",
    "Botões de menu (•••) têm aria-label e estado anchored para realce visual.",
    "Componente de seleção de pessoa (PessoaSelector) usa combobox com busca.",
])

doc.add_page_break()

# ───────────────────────── 7. ESTADO E VALIDAÇÕES ─────────────────────────
h1("7. Estado, derivações e validações")

h3("7.1 Estado da view (resumo)")
bullets([
    "Cabeçalho: data, safra (derivada), proprietario, fazenda, retiro, local",
    "Modo: fromId: boolean (toggle Brinco) ; lrExpanded: boolean (tela cheia).",
    "Coletivo: cats: NascCat[] (declarados) ; catSel: string ; totalStr: string.",
    "Individual: detalhe: NascDetalhe[] ; entryValues: Record<id,string> (linha de inclusão).",
    "Sanitário: sanItems: SanItem[] ; sanOpen: boolean.",
    "Persistência: movimentos: MovimentoNasc[] ; editingId: string|null ; atribuirTargetId: string|null.",
    "Aba: aba: 'lancar' | 'historico'.",
])

h3("7.2 Derivações principais (funções puras em util.ts)")
bullets([
    "safraDaData(iso): retorna 'YYYY/YYYY' (jul→jun).",
    "proximoApelido(prev): incrementa o número preservando prefixo/sufixo/zeros.",
    "tallyPorCategoria(detalhe): {catId: count} dos animais detalhados.",
    "somaCategorias(cats): soma das quantidades declaradas.",
    "statusFrom(naoIdent): 'pendente' se >0; senão 'conciliado'.",
    "sexoFromCategoria(categories, catId): normaliza para 'Macho' | 'Fêmea'.",
    "consolidated: junta cats + tallyPorCategoria(detalhe) por catId — alimenta o CategoriaGrid.",
])

h3("7.3 Validações antes do Save")
bullets([
    "organizationId presente (selectedOrganization).",
    "Soma de declarados + detalhados > 0.",
    "Cada animal detalhado precisa de apelido e categoria.",
    "Categorias do select são SEMPRE filtradas por grupo='bezerros_mamando'.",
    "Data válida (YYYY-MM-DD) — validada também no servidor.",
])

doc.add_page_break()

# ───────────────────────── 8. CASOS DE USO ─────────────────────────
h1("8. Casos de uso (cenários)")

h3("CU-01 — Lançamento coletivo simples (capataz no celular)")
para("Maria, capataz, registra 18 nascimentos do dia em 30 segundos: seleciona Data, "
     "confirma Fazenda/Retiro/Local pré-preenchidos, marca 18 cab. de 'Bezerros Mamando', "
     "Salvar. Lançamento entra 'pendente' (18 não identificados) e aparece na aba Registros.")

h3("CU-02 — Atribuição posterior pelo técnico (escritório)")
para("José, técnico, abre a aba Registros, expande o lançamento de Maria, clica ••• → "
     "Atribuir ID. Inclui os 18 bezerros um a um (ID Manejo, RFID, Peso). A cada Add, "
     "naoIdentificados decrementa. Ao terminar, status muda para 'conciliado' automaticamente.")

h3("CU-03 — Lançamento em massa por planilha")
para("Carlos, zootecnista, exporta o modelo .xlsx (com os mesmos campos da tela), preenche "
     "no Excel com 80 bezerros, importa. O modal de validação aponta 2 categorias inexistentes "
     "e 1 duplicado. Carlos corrige no Excel, importa de novo. Salvar. 80 fichas viram nascimento "
     "em um único lançamento conciliado.")

h3("CU-04 — Configuração inicial por organização")
para("No primeiro uso, o administrador abre 'Configurar campos': move 'SISBOV' para "
     "'Desativado' (fazenda não usa SISBOV), promove 'Pelagem' para 'Linha Superior' e ativa "
     "Numeração Automática. A configuração vale para TODOS os usuários da organização.")

h3("CU-05 — Edição de um lançamento")
para("Foi descoberto que o lançamento do dia 20/06 tinha 1 bezerro a mais. Na aba Registros, "
     "••• → Editar. O lançamento é reaberto NO TOPO do formulário (banner laranja avisando). "
     "O usuário remove o bezerro, clica 'Salvar alterações'. O PUT substitui as fichas pelo "
     "novo conjunto. Status é reavaliado.")

doc.add_page_break()

# ───────────────────────── 9. EDGE CASES ─────────────────────────
h1("9. Casos de borda e tratamento de erro")

h3("9.1 Comportamento esperado")
bullets([
    "Sem organização selecionada → bloqueia POST/PUT/DELETE com toast 'Selecione uma organização antes de salvar'.",
    "Categoria com sexo cru ('M'/'F'/'macho'/'feminino') é normalizada para 'Macho'/'Fêmea' (sexoFromCategoria).",
    "Trocar de Retiro reseta o Local (evita Local órfão).",
    "Raça do registry: se NÃO há raças cadastradas, usa lista estática (RACAS).",
    "ID Manejo duplicado dentro do MESMO lançamento → erro visível na validação da importação.",
    "Importação > 1MB ou .xlsx corrompido → toast 'Erro ao ler a planilha. Verifique se é .xlsx, .xls ou .csv.'",
])

h3("9.2 Falhas de rede")
bullets([
    "Listagem (GET) com erro → toast 'Erro ao carregar nascimentos'.",
    "Save (POST/PUT) com erro → mantém o formulário, mostra toast com mensagem do servidor (campo error).",
    "Delete com erro → não remove da lista, toast com motivo.",
])

h3("9.3 Concorrência (multi-usuário)")
bullets([
    "Não há lock no movimento — a UI confia em última-escrita-vence.",
    "Atribuir ID via POST add-ficha é atômico no servidor (decremento + recálculo).",
    "Para Editar, a UI substitui o conjunto de fichas; reentrada simultânea de dois usuários pode sobrescrever — aceitável para o domínio.",
])

doc.add_page_break()

# ───────────────────────── 10. INSTRUÇÕES P/ DEV ─────────────────────────
h1("10. Roadmap de implementação (instruções para o time)")

h3("10.1 Ordem sugerida de entrega")
num_list([
    "Backend — schema e endpoints: criar tabelas e o handler /api/nascimentos (CRUD + add-ficha) + /api/nascimento-field-config (GET/PUT).",
    "UI base (modo COLETIVO): cabeçalho + cartão master + painel direito de Distribuição. Salvar/Editar/Excluir.",
    "Configuração de campos: modal 'Configurar campos' + persistência por organização. Reset para defaults.",
    "Modo INDIVIDUAL: 'Defina seus campos' (FichaInclusaoForm + DefinaCamposPanel). Inclusão de animais com sugestão de próximo ID.",
    "Atribuir ID (painel posterior + add-ficha do servidor) — chamado a partir da aba Registros e da menu ••• do lançamento.",
    "Aba Registros (master-detail) com régua arrastável e filtro por categoria expandida.",
    "Sanitário (recolhível) — listas estáticas de medicamentos/protocolos.",
    "Importação por planilha (exportLancamentoTemplate + ImportarPlanilhaModal).",
    "Tela cheia (FullscreenLancamento) e atalhos (Esc).",
    "Acessibilidade, responsividade, microinterações.",
])

h3("10.2 Dependências externas")
bullets([
    "PostgreSQL (Neon) + Drizzle ORM (push manual via tmp/create-*-table.ts para tabelas aditivas).",
    "@dnd-kit (sortable) — modal de configuração de campos.",
    "lucide-react — todos os ícones (Plus, Save, Check, Info, List, Tags, ChevronDown, Pencil, Trash2, MoreHorizontal, etc.).",
    "xlsx — leitura/escrita da planilha modelo (já no projeto).",
    "Tailwind v4 + container queries — layout responsivo.",
    "better-auth — cookie/bearer no /api/nascimentos.",
])

h3("10.3 Pontos de atenção (anti-patterns que JÁ ocorreram)")
bullets([
    "Rota /api/nascimentos PRECISA ser registrada à mão em server-dev.ts (vide memória project_dev_api_routes.md). Sem isso, o dev retorna HTML 404.",
    "Tabelas aditivas: drizzle-kit push tem tablesFilter (allowlist). Criar via tmp/create-nascimentos-table.ts (tsx).",
    "Sexo: NUNCA pedir ao usuário no modo nascimento — é sempre derivado da categoria.",
    "Categoria atual do animal é a UNIÃO de nascimento_fichas + fichas_animal (dedup por apelido) — vide animalRegistry.ts.",
    "Não escrever Sanitário como tabela ainda; mantém-se como snapshot jsonb.",
    "Ao implementar a tela mobile, manter a régua interna do cartão como empilhamento (container queries).",
])

h3("10.4 Critérios de aceite")
bullets([
    "Lançar 50 nascimentos coletivos em < 30s (sem rede lenta).",
    "Atribuir ID a 20 bezerros em < 2 min via teclado (Enter envia).",
    "Importar planilha de 100 linhas com erros mistos: o modal mostra a lista detalhada; só as válidas são confirmadas.",
    "Editar/Excluir mantêm a UI coerente (sem fichas órfãs, sem status incoerente).",
    "Configuração de campos persiste e reaparece após F5 com a mesma ordem/destinos.",
])

doc.add_page_break()

# ───────────────────────── 11. DICIONÁRIO ─────────────────────────
h1("11. Dicionário de termos")
glossary = [
    ["Termo", "Significado"],
    ["ID Manejo (apelido)", "Identificador interno do animal no manejo (texto livre). Único por organização."],
    ["SISBOV",   "Número do Sistema de Identificação Brasileiro (cadeia de carne e leite)."],
    ["RFID",     "Brinco eletrônico (radiofrequência)."],
    ["Safra",    "Ano-safra (jul→jun) — calculado pela data do lançamento."],
    ["Retiro",   "Subdivisão da fazenda (ex.: Sede, Pirituba, Brejo). Nome livre, não-FK."],
    ["Local",    "Folha do nível de áreas (pasto, curral, piquete) — FK farm_locais."],
    ["catDecl",  "Distribuição CONSOLIDADA por categoria (declarado + detalhado)."],
    ["naoIdentificados", "Quantidade ainda pendente de individualização (sem ficha)."],
    ["Conciliado", "Status quando naoIdentificados == 0."],
    ["Pendente",   "Status quando naoIdentificados > 0."],
    ["Kit 'Defina seus campos'", "Configuração compartilhada com Compra/Venda/Morte. Cada campo tem destino e ordem."],
    ["Lançamento Rápido", "Sinônimo do modo individual (linha de inclusão por animal)."],
    ["Master-detail", "Padrão da aba Registros: lista em cima, fichas embaixo, régua arrastável."],
]
add_table(glossary, col_widths=[5.0, 22.0])

doc.add_page_break()

# ───────────────────────── 12. APÊNDICE ─────────────────────────
h1("12. Apêndice — Mapeamento código ↔ funcionalidade")
para("Para auxiliar o time durante a reescrita, abaixo está a referência cruzada com a "
     "implementação atual em React/TypeScript.", gray=True, italic=True)

xref = [
    ["Funcionalidade",                       "Arquivo principal"],
    ["Tela principal e orquestração",        "agents/pecuario/nascimento/NascimentoView.tsx"],
    ["Catálogo de campos / defaults",        "agents/pecuario/nascimento/fieldRegistry.ts"],
    ["Tipos do domínio",                     "agents/pecuario/nascimento/types.ts"],
    ["Derivações (safra, próximo ID, tally)","agents/pecuario/nascimento/util.ts"],
    ["Grid 'Distribuição por categoria'",    "agents/pecuario/nascimento/CategoriaGrid.tsx"],
    ["Aba Registros (master-detail)",        "agents/pecuario/nascimento/LancamentosRecentes.tsx"],
    ["Painel 'Atribuir ID'",                 "agents/pecuario/nascimento/AtribuirIdPanel.tsx"],
    ["Seção Sanitário",                      "agents/pecuario/nascimento/SanitarioSection.tsx"],
    ["Kit 'Defina seus campos'",             "agents/pecuario/fichas/DefinaCamposPanel.tsx"],
    ["Modal de configuração de campos",      "agents/pecuario/fichas/CamposConfigModal.tsx"],
    ["Shell tela cheia",                     "agents/pecuario/fichas/FullscreenLancamento.tsx"],
    ["Form de inclusão",                     "agents/pecuario/fichas/FichaInclusaoForm.tsx"],
    ["Importação de planilha",               "agents/pecuario/fichas/ImportarPlanilhaModal.tsx + importTemplate.ts"],
    ["Cliente HTTP /api/nascimentos",        "lib/api/nascimentosClient.ts"],
    ["Repositório (Drizzle)",                "src/DB/repositories/nascimentos.ts"],
    ["Handler /api/nascimentos",             "api/nascimentos.ts"],
    ["Schema (tabelas)",                     "src/DB/schema.ts (linhas 1480–1536)"],
    ["Field config (cliente)",               "lib/api/nascimentoFieldConfigClient.ts"],
    ["Field config (repo)",                  "src/DB/repositories/nascimentoFieldConfig.ts"],
    ["Field config (handler)",               "api/nascimento-field-config.ts"],
]
add_table(xref, col_widths=[7.5, 20.0])

doc.add_paragraph()
hr_para()
para("— FIM —", italic=True, gray=True)

# Salvar
doc.save(OUT_FILE)
print("OK →", OUT_FILE)
